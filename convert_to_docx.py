import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_styled_docx(md_path, docx_path, title_text):
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title
    title = doc.add_heading(level=0)
    run = title.add_run(title_text)
    run.font.name = 'Helvetica'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(19, 39, 67) # Dark Navy
    
    doc.add_paragraph() # Spacing
    
    if os.path.exists(md_path):
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            if line_str.startswith("# "):
                h = doc.add_heading(level=1)
                r = h.add_run(line_str[2:])
                r.font.color.rgb = RGBColor(19, 39, 67)
                r.font.size = Pt(18)
            elif line_str.startswith("## "):
                h = doc.add_heading(level=2)
                r = h.add_run(line_str[3:])
                r.font.color.rgb = RGBColor(163, 230, 53) # DARWIN Lime Green accent
                r.font.size = Pt(14)
            elif line_str.startswith("### "):
                h = doc.add_heading(level=3)
                r = h.add_run(line_str[4:])
                r.font.color.rgb = RGBColor(56, 68, 77)
                r.font.size = Pt(12)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line_str[2:])
            elif line_str.startswith("```"):
                p = doc.add_paragraph()
                r = p.add_run(line_str)
                r.font.name = 'Courier New'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(40, 40, 40)
            else:
                p = doc.add_paragraph()
                p.add_run(line_str)
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")

out_dir = os.path.expanduser("~/Darwin/DOCX_Documentation")
os.makedirs(out_dir, exist_ok=True)

artifacts_dir = "/Users/joelagyeman/.gemini/antigravity-ide/brain/3c780013-73d8-4b2f-9603-3bbd4cfa3076"

mappings = [
    (os.path.join(artifacts_dir, "project_architecture.md"), os.path.join(out_dir, "01_DARWIN_Architecture.docx"), "DARWIN Sovereign Capability Engine Architecture"),
    (os.path.join(artifacts_dir, "walkthrough.md"), os.path.join(out_dir, "02_DARWIN_Walkthrough.docx"), "DARWIN Monorepo & Engine Walkthrough"),
    (os.path.join(artifacts_dir, "model_decision_matrix.md"), os.path.join(out_dir, "03_C++_Structural_Entropy_Router.docx"), "C++ Structural Entropy AI Model Router"),
    (os.path.join(artifacts_dir, "security_assessment.md"), os.path.join(out_dir, "04_Security_Assessment_Akan_Ontology.docx"), "DARWIN Security Assessment & Akan 4D Ontology"),
    (os.path.join(artifacts_dir, "cybersec_swot.md"), os.path.join(out_dir, "05_Cybersecurity_SWOT_Analysis.docx"), "DARWIN Cybersecurity SWOT Analysis"),
    (os.path.join(artifacts_dir, "AFRICAN_INDUSTRIALISATION.md"), os.path.join(out_dir, "06_African_Industrialisation_Strategy.docx"), "African Industrialisation & Sovereign Engine Strategy"),
    (os.path.expanduser("~/Darwin/boot_darwin.sh"), os.path.join(out_dir, "07_DARWIN_Master_Boot_Script.docx"), "DARWIN Master Launcher Script Specifications")
]

for src, dst, title in mappings:
    create_styled_docx(src, dst, title)
