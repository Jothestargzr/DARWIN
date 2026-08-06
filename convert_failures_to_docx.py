import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor

def code_file_to_docx(src_path, docx_path, title_text):
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Title
    title = doc.add_heading(level=0)
    run = title.add_run(title_text)
    run.font.name = 'Helvetica'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(19, 39, 67) # Dark Navy
    
    doc.add_paragraph() # Spacing
    
    if os.path.exists(src_path):
        with open(src_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
            
        p = doc.add_paragraph()
        for idx, line in enumerate(lines, 1):
            r_num = p.add_run(f"{idx:4d}: ")
            r_num.font.name = 'Courier New'
            r_num.font.size = Pt(9)
            r_num.font.color.rgb = RGBColor(120, 120, 120)
            
            r_code = p.add_run(line)
            r_code.font.name = 'Courier New'
            r_code.font.size = Pt(9)
            r_code.font.color.rgb = RGBColor(30, 30, 30)
            
    doc.save(docx_path)
    print(f"Created DOCX: {docx_path}")

out_dir = os.path.expanduser("~/Darwin/DOCX_Documentation")
os.makedirs(out_dir, exist_ok=True)

failure_files = [
    (
        os.path.expanduser("~/Darwin/.github/workflows/build-darwin-os.yml"),
        os.path.join(out_dir, "FAIL_01_build-darwin-os_workflow.docx"),
        "Cloud Compiler Workflow Specification (build-darwin-os.yml)"
    ),
    (
        os.path.expanduser("~/Darwin/BrowserOS/packages/browseros/bos_build/browseros.py"),
        os.path.join(out_dir, "FAIL_02_browseros_entrypoint.docx"),
        "bos_build Main Entrypoint CLI (browseros.py)"
    ),
    (
        os.path.expanduser("~/Darwin/BrowserOS/packages/browseros/bos_build/cli/dev.py"),
        os.path.join(out_dir, "FAIL_03_dev_cli_commands.docx"),
        "bos_build Dev Subcommand CLI (dev.py)"
    ),
    (
        os.path.expanduser("~/Darwin/BrowserOS/packages/browseros/bos_build/steps/storage/download.py"),
        os.path.join(out_dir, "FAIL_04_download_resources_step.docx"),
        "bos_build R2 Download Resources Step (download.py)"
    )
]

for src, dst, title in failure_files:
    code_file_to_docx(src, dst, title)
