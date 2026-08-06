import os
import sys
import time
from concurrent.futures import ProcessPoolExecutor
from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT_DIR = os.path.abspath(os.path.dirname(__file__))
OUTPUT_DIR = os.path.join(ROOT_DIR, "DARWIN_DOCX_COPIES")

IGNORE_DIRS = {
    '.git', 'node_modules', '__pycache__', '.venv', 'venv', 'out', 'dist', 
    'build', '.next', '.wxt', 'DARWIN_DOCX_COPIES', 'DOCX_Documentation', 
    'coverage', '.turbo', '.cache', 'target', '.output'
}

IGNORE_EXTS = {
    '.docx', '.dmg', '.png', '.jpg', '.jpeg', '.gif', '.ico', '.pdf', '.zip', 
    '.tar', '.gz', '.7z', '.exe', '.bin', '.pyc', '.ds_store', '.ttf', '.woff', 
    '.woff2', '.eot', '.mp3', '.mp4', '.mov', '.avi', '.icns', '.o', '.a', 
    '.so', '.dylib', '.wasm', '.lock'
}

def convert_single_file(args):
    src_abs_path, rel_path, dst_docx_path = args
    try:
        os.makedirs(os.path.dirname(dst_docx_path), exist_ok=True)
        
        doc = Document()
        
        # Margins
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            
        # Title
        title_p = doc.add_heading(level=1)
        r_title = title_p.add_run(f"DARWIN Source: {rel_path}")
        r_title.font.name = 'Helvetica'
        r_title.font.size = Pt(16)
        r_title.font.bold = True
        r_title.font.color.rgb = RGBColor(19, 39, 67) # Dark Navy
        
        # Metadata
        meta_p = doc.add_paragraph()
        r_meta = meta_p.add_run(f"File Path: {rel_path}\nOriginal File: {os.path.basename(src_abs_path)}")
        r_meta.font.name = 'Helvetica'
        r_meta.font.size = Pt(9.5)
        r_meta.font.italic = True
        r_meta.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph() # Spacing
        
        # Read content
        with open(src_abs_path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
            
        p = doc.add_paragraph()
        for idx, line in enumerate(lines, 1):
            # Sanitize control characters for openxml
            sanitized_line = "".join(ch for ch in line if ch == '\n' or ch == '\r' or ch == '\t' or (ord(ch) >= 32 and ord(ch) != 127))
            
            r_num = p.add_run(f"{idx:5d}  ")
            r_num.font.name = 'Courier New'
            r_num.font.size = Pt(8.5)
            r_num.font.color.rgb = RGBColor(140, 140, 140)
            
            r_code = p.add_run(sanitized_line)
            r_code.font.name = 'Courier New'
            r_code.font.size = Pt(9.0)
            r_code.font.color.rgb = RGBColor(30, 30, 30)
            
        doc.save(dst_docx_path)
        return True, rel_path, None
    except Exception as e:
        return False, rel_path, str(e)

def main():
    start_time = time.time()
    print(f"🚀 Scanning DARWIN repository at: {ROOT_DIR}")
    
    file_tasks = []
    
    for dirpath, dirnames, filenames in os.walk(ROOT_DIR):
        dirnames[:] = [d for d in dirnames if d not in IGNORE_DIRS and not d.startswith('.git')]
        for f in filenames:
            if f.startswith('.') and f != '.gitignore' and f != '.env':
                continue
            ext = os.path.splitext(f)[1].lower()
            if ext in IGNORE_EXTS:
                continue
                
            src_abs_path = os.path.normpath(os.path.join(dirpath, f))
            rel_path = os.path.relpath(src_abs_path, ROOT_DIR)
            
            # Skip output folder itself
            if rel_path.startswith("DARWIN_DOCX_COPIES"):
                continue
                
            # Check if text file by inspecting binary bytes
            try:
                with open(src_abs_path, 'rb') as fp:
                    chunk = fp.read(1024)
                    if b'\x00' in chunk:
                        continue
            except Exception:
                continue
                
            dst_docx_path = os.path.join(OUTPUT_DIR, rel_path + ".docx")
            file_tasks.append((src_abs_path, rel_path, dst_docx_path))
            
    total_files = len(file_tasks)
    print(f"📋 Found {total_files} text files to convert into DOCX.")
    print(f"📁 Output Directory: {OUTPUT_DIR}")
    print("⏳ Starting parallel conversion across CPU cores...")
    
    success_count = 0
    fail_count = 0
    
    # Process in parallel using ProcessPoolExecutor
    workers = min(os.cpu_count() or 4, 16)
    with ProcessPoolExecutor(max_workers=workers) as executor:
        results = executor.map(convert_single_file, file_tasks, chunksize=20)
        
        for idx, (success, rel_path, err) in enumerate(results, 1):
            if success:
                success_count += 1
            else:
                fail_count += 1
                print(f"❌ Error converting {rel_path}: {err}")
                
            if idx % 250 == 0 or idx == total_files:
                elapsed = time.time() - start_time
                print(f"Progress: [{idx}/{total_files}] files converted ({idx/total_files*100:.1f}%) in {elapsed:.1f}s")
                
    elapsed_total = time.time() - start_time
    print("\n==========================================")
    print(f"🎉 CONVERSION COMPLETE!")
    print(f"✅ Successfully converted: {success_count} files")
    if fail_count > 0:
        print(f"⚠️ Failed: {fail_count} files")
    print(f"⏱️ Total Execution Time: {elapsed_total:.2f} seconds")
    print(f"📂 DOCX Copies location: {OUTPUT_DIR}")
    print("==========================================")

if __name__ == "__main__":
    main()
